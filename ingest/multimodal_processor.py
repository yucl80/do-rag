from typing import Dict, List, Tuple, Any, Optional, Callable
import pdfplumber
import docx
import pandas as pd
from PIL import Image
import pytesseract
import cv2
import numpy as np
from transformers import AutoProcessor, AutoModel
import torch
from sentence_transformers import SentenceTransformer
import spacy
import networkx as nx
from pathlib import Path
import logging
from tqdm import tqdm
import gc

class MultimodalProcessor:
    """Handles processing of various document types and multimodal content."""
    
    def __init__(self, 
                 text_model_name: str = 'all-MiniLM-L6-v2',
                 vision_model_name: str = 'google/vit-base-patch16-224',
                 device: str = 'cuda' if torch.cuda.is_available() else 'cpu',
                 chunk_size: int = 500,
                 max_image_size: Tuple[int, int] = (1024, 1024)):
        """
        Initialize the multimodal processor.
        
        Args:
            text_model_name: Name of the text embedding model
            vision_model_name: Name of the vision model for image processing
            device: Device to run models on
            chunk_size: Maximum size of text chunks
            max_image_size: Maximum size for image processing
        """
        self.device = device
        self.chunk_size = chunk_size
        self.max_image_size = max_image_size
        
        try:
            self.text_model = SentenceTransformer(text_model_name)
            self.vision_processor = AutoProcessor.from_pretrained(vision_model_name)
            self.vision_model = AutoModel.from_pretrained(vision_model_name).to(device)
            self.nlp = spacy.load("en_core_web_sm")
        except Exception as e:
            raise RuntimeError(f"Failed to initialize models: {str(e)}")
        
        # Initialize document type handlers
        self.handlers = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.xlsx': self._process_excel,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image,
            '.txt': self._process_text
        }
        
        # Setup logging
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)
        
    def process_document(self, file_path: str, progress_callback: Optional[Callable[[float], None]] = None) -> Dict[str, Any]:
        """
        Process a document and extract its content.
        
        Args:
            file_path: Path to the document
            progress_callback: Optional callback function for progress updates
            
        Returns:
            Dictionary containing processed content and metadata
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file type is not supported
            RuntimeError: If processing fails
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        handler = self.handlers.get(file_path.suffix.lower())
        if not handler:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
            
        try:
            if progress_callback:
                progress_callback(0.0)
            content = handler(file_path, progress_callback)
            if progress_callback:
                progress_callback(1.0)
            return content
        except Exception as e:
            self.logger.error(f"Error processing document {file_path}: {str(e)}")
            raise RuntimeError(f"Failed to process document: {str(e)}")
        
    def _process_pdf(self, file_path: Path, progress_callback: Optional[Callable[[float], None]] = None) -> Dict[str, Any]:
        """Process PDF documents."""
        content = {
            'text_chunks': [],
            'tables': [],
            'images': [],
            'metadata': {}
        }
        
        try:
            with pdfplumber.open(file_path) as pdf:
                content['metadata']['pages'] = len(pdf.pages)
                total_pages = len(pdf.pages)
                
                for page_num, page in enumerate(tqdm(pdf.pages, desc="Processing PDF pages")):
                    try:
                        # Extract text
                        text = page.extract_text()
                        if text:
                            chunks = self._create_semantic_chunks(text)
                            content['text_chunks'].extend(chunks)
                        
                        # Extract tables
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                try:
                                    df = pd.DataFrame(table[1:], columns=table[0])
                                    content['tables'].append({
                                        'page': page_num + 1,
                                        'data': df.to_dict('records')
                                    })
                                except Exception as e:
                                    self.logger.warning(f"Failed to process table on page {page_num + 1}: {str(e)}")
                        
                        # Extract images
                        images = page.images
                        if images:
                            for img in images:
                                content['images'].append({
                                    'page': page_num + 1,
                                    'bbox': img['bbox'],
                                    'width': img['width'],
                                    'height': img['height']
                                })
                        
                        if progress_callback:
                            progress_callback((page_num + 1) / total_pages)
                            
                    except Exception as e:
                        self.logger.warning(f"Failed to process page {page_num + 1}: {str(e)}")
                        continue
                        
        except Exception as e:
            raise RuntimeError(f"Failed to process PDF: {str(e)}")
            
        return content
        
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Process Word documents."""
        content = {
            'text_chunks': [],
            'tables': [],
            'images': [],
            'metadata': {}
        }
        
        doc = docx.Document(file_path)
        content['metadata']['paragraphs'] = len(doc.paragraphs)
        
        # Process paragraphs
        current_chunk = []
        for para in doc.paragraphs:
            if para.text.strip():
                current_chunk.append(para.text)
                if len(' '.join(current_chunk)) > 500:  # Chunk size threshold
                    content['text_chunks'].append(' '.join(current_chunk))
                    current_chunk = []
        
        if current_chunk:
            content['text_chunks'].append(' '.join(current_chunk))
        
        # Process tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            if table_data:
                df = pd.DataFrame(table_data[1:], columns=table_data[0])
                content['tables'].append({
                    'data': df.to_dict('records')
                })
        
        # Process images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                content['images'].append({
                    'type': 'embedded',
                    'target': rel.target_ref
                })
        
        return content
        
    def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Process Excel documents."""
        content = {
            'tables': [],
            'metadata': {}
        }
        
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        content['metadata']['sheets'] = excel_file.sheet_names
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            content['tables'].append({
                'sheet': sheet_name,
                'data': df.to_dict('records')
            })
        
        return content
        
    def _process_image(self, file_path: Path, progress_callback: Optional[Callable[[float], None]] = None) -> Dict[str, Any]:
        """Process image files."""
        content = {
            'text': [],
            'features': None,
            'metadata': {}
        }
        
        try:
            # Load and resize image
            image = Image.open(file_path)
            content['metadata']['original_size'] = image.size
            content['metadata']['mode'] = image.mode
            
            # Resize if too large
            if image.size[0] > self.max_image_size[0] or image.size[1] > self.max_image_size[1]:
                image.thumbnail(self.max_image_size, Image.Resampling.LANCZOS)
                content['metadata']['processed_size'] = image.size
            
            if progress_callback:
                progress_callback(0.3)
            
            # Extract text using OCR
            try:
                text = pytesseract.image_to_string(image)
                if text.strip():
                    content['text'] = self._create_semantic_chunks(text)
            except Exception as e:
                self.logger.warning(f"OCR failed: {str(e)}")
            
            if progress_callback:
                progress_callback(0.6)
            
            # Extract image features with memory management
            try:
                inputs = self.vision_processor(images=image, return_tensors="pt").to(self.device)
                with torch.no_grad():
                    features = self.vision_model(**inputs).last_hidden_state.mean(dim=1)
                content['features'] = features.cpu().numpy()
                
                # Clear GPU memory
                del inputs
                del features
                torch.cuda.empty_cache()
                gc.collect()
            except Exception as e:
                self.logger.warning(f"Feature extraction failed: {str(e)}")
            
            if progress_callback:
                progress_callback(1.0)
                
        except Exception as e:
            raise RuntimeError(f"Failed to process image: {str(e)}")
            
        return content
        
    def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """Process plain text files."""
        content = {
            'text_chunks': [],
            'metadata': {}
        }
        
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            
        content['text_chunks'] = self._create_semantic_chunks(text)
        content['metadata']['size'] = len(text)
        
        return content
        
    def _create_semantic_chunks(self, text: str) -> List[str]:
        """Create semantic chunks from text using improved strategy."""
        doc = self.nlp(text)
        chunks = []
        current_chunk = []
        current_length = 0
        
        # Track paragraph boundaries
        paragraph_boundaries = []
        for token in doc:
            if token.is_space and '\n' in token.text:
                paragraph_boundaries.append(token.i)
        
        for sent in doc.sents:
            sent_text = sent.text.strip()
            if not sent_text:
                continue
                
            sent_length = len(sent_text)
            
            # Check if sentence crosses paragraph boundary
            crosses_boundary = any(token.i in paragraph_boundaries for token in sent)
            
            # Start new chunk if:
            # 1. Current would exceed limit
            # 2. Sentence crosses paragraph boundary
            # 3. Current chunk is not empty and sentence starts with a new paragraph
            if (current_length + sent_length > self.chunk_size or 
                crosses_boundary or 
                (current_chunk and sent.start_char > current_chunk[-1].end_char + 1)):
                
                if current_chunk:
                    chunks.append(' '.join(current_chunk))
                current_chunk = [sent_text]
                current_length = sent_length
            else:
                current_chunk.append(sent_text)
                current_length += sent_length
        
        # Add remaining chunk
        if current_chunk:
            chunks.append(' '.join(current_chunk))
            
        return chunks
        
    def get_embeddings(self, content: Dict[str, Any]) -> Dict[str, np.ndarray]:
        """Generate embeddings for different types of content."""
        embeddings = {}
        
        try:
            # Generate text embeddings
            if 'text_chunks' in content:
                text_chunks = content['text_chunks']
                if text_chunks:
                    embeddings['text'] = self.text_model.encode(text_chunks)
            
            # Generate image embeddings
            if 'features' in content and content['features'] is not None:
                embeddings['image'] = content['features']
                
        except Exception as e:
            self.logger.error(f"Failed to generate embeddings: {str(e)}")
            raise RuntimeError(f"Embedding generation failed: {str(e)}")
            
        return embeddings 